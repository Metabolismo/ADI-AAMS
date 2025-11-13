#!/usr/bin/env python3
"""
LaTeX to GATM DOCX Converter - Versione Corretta
Converte un abstract LaTeX nel formato DOCX richiesto dal GATM
"""

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE

class FixedLatexToGATMConverter:
    def __init__(self):
        self.doc = Document()
        self.setup_document_style()

    def setup_document_style(self):
        """Configura gli stili del documento secondo le specifiche GATM"""
        # Margini: 2.5cm
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.98)    # ~2.5cm
            section.bottom_margin = Inches(0.98)
            section.left_margin = Inches(0.98)
            section.right_margin = Inches(0.98)

        # Stile Normal per il corpo del testo
        styles = self.doc.styles
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(11)

        # Formattazione paragrafo
        normal_paragraph = normal_style.paragraph_format
        normal_paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal_paragraph.line_spacing = 1.5
        normal_paragraph.first_line_indent = Inches(0.197)  # 0.5cm
        normal_paragraph.space_after = Pt(0)

    def extract_latex_content(self, latex_text):
        """Estrae i componenti dall'abstract LaTeX usando pattern specifici"""
        content = {}

        # Estrai titolo - pattern più robusto
        title_pattern = r'\\abstracttitle\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}'
        title_match = re.search(title_pattern, latex_text)
        if title_match:
            content['title'] = self.clean_title_text(title_match.group(1))
        else:
            content['title'] = ""

        # Estrai autori
        authors_pattern = r'\\abstractauthors\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}'
        authors_match = re.search(authors_pattern, latex_text)
        if authors_match:
            content['authors'] = authors_match.group(1)
        else:
            content['authors'] = ""

        # Estrai affiliazioni
        affil_pattern = r'\\abstractaffiliations\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}'
        affil_match = re.search(affil_pattern, latex_text)
        if affil_match:
            content['affiliations'] = affil_match.group(1)
        else:
            content['affiliations'] = ""

        # Estrai corpo del testo - dal comando affiliations alla bibliografia
        body_start = re.search(r'\\abstractaffiliations\{[^}]*\}', latex_text)
        body_end = re.search(r'\\gatabase', latex_text)

        if body_start and body_end:
            body_text = latex_text[body_start.end():body_end.start()]
            content['body'] = self.clean_body_text(body_text)
        else:
            content['body'] = ""

        # Estrai bibliografia - da \gatabase a \end{document}
        bib_start = re.search(r'\\gatabase', latex_text)
        bib_end = re.search(r'\\end\{document\}', latex_text)

        if bib_start and bib_end:
            bib_text = latex_text[bib_start.end():bib_end.start()]
            content['bibliography'] = self.extract_bibliography_entries(bib_text)
        else:
            content['bibliography'] = []

        return content

    def clean_title_text(self, text):
        """Pulisce il testo del titolo"""
        # Converti corsivi LaTeX in markdown temporaneo
        text = re.sub(r'\\emph\{([^}]+)\}', r'*\1*', text)
        text = re.sub(r'\\textit\{([^}]+)\}', r'*\1*', text)

        # Rimuovi altri comandi LaTeX
        text = re.sub(r'\\[a-zA-Z]+\*?(\[[^\]]*\])?\{([^}]*)\}', r'\2', text)
        text = re.sub(r'\\[a-zA-Z]+\*?', '', text)

        # Pulisci spazi
        text = re.sub(r'\s+', ' ', text).strip()
        return text

    def clean_body_text(self, text):
        """Pulisce il corpo del testo"""
        # Rimuovi commenti LaTeX
        text = re.sub(r'%.*$', '', text, flags=re.MULTILINE)

        # Converti corsivi
        text = re.sub(r'\\emph\{([^}]+)\}', r'*\1*', text)
        text = re.sub(r'\\textit\{([^}]+)\}', r'*\1*', text)

        # Gestisci citazioni nel formato (Autore, Anno)
        text = re.sub(r'\(([^)]+),\s*(\d{4})\)', r'(\1, \2)', text)

        # Rimuovi altri comandi LaTeX
        text = re.sub(r'\\[a-zA-Z]+\*?(\[[^\]]*\])?\{([^}]*)\}', r'\2', text)
        text = re.sub(r'\\[a-zA-Z]+\*?', '', text)

        # Pulisci spazi multipli e newline
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Mantieni paragrafi
        text = re.sub(r'[ \t]+', ' ', text)  # Rimuovi spazi multipli
        text = text.strip()

        return text

    def extract_bibliography_entries(self, bib_text):
        """Estrae le voci bibliografiche"""
        entries = []

        # Pulisci il testo della bibliografia
        bib_text = re.sub(r'%.*$', '', bib_text, flags=re.MULTILINE)
        bib_text = bib_text.strip()

        # Dividi per linee e raggruppa le voci
        lines = [line.strip() for line in bib_text.split('\n') if line.strip()]

        current_entry = ""
        for line in lines:
            # Se la linea inizia con una maiuscola e abbiamo già un'entry, inizia una nuova voce
            if line and line[0].isupper() and current_entry:
                entries.append(self.clean_bibliography_entry(current_entry))
                current_entry = line
            else:
                current_entry += " " + line if current_entry else line

        # Aggiungi l'ultima voce
        if current_entry:
            entries.append(self.clean_bibliography_entry(current_entry))

        return entries

    def clean_bibliography_entry(self, entry):
        """Pulisce una singola voce bibliografica"""
        # Converti corsivi
        entry = re.sub(r'\\emph\{([^}]+)\}', r'*\1*', entry)
        entry = re.sub(r'\\textit\{([^}]+)\}', r'*\1*', entry)

        # Converti en-dash LaTeX
        entry = re.sub(r'--', '–', entry)

        # Rimuovi altri comandi LaTeX
        entry = re.sub(r'\\[a-zA-Z]+\*?(\[[^\]]*\])?\{([^}]*)\}', r'\2', entry)
        entry = re.sub(r'\\[a-zA-Z]+\*?', '', entry)

        # Pulisci spazi
        entry = re.sub(r'\s+', ' ', entry).strip()

        return entry

    def add_title(self, title):
        """Aggiunge il titolo (16pt, grassetto, centrato)"""
        title_paragraph = self.doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.paragraph_format.space_after = Pt(6)
        title_paragraph.paragraph_format.first_line_indent = Inches(0)  # Nessuna indentazione per il titolo

        # Processa il titolo con corsivi
        self.add_formatted_text(title_paragraph, title, font_size=16, bold=True)

    def add_authors(self, authors):
        """Aggiunge gli autori (14pt, centrato)"""
        authors_paragraph = self.doc.add_paragraph()
        authors_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        authors_paragraph.paragraph_format.space_after = Pt(3)
        authors_paragraph.paragraph_format.first_line_indent = Inches(0)

        # Gestisci gli apici per le affiliazioni
        parts = re.split(r'(\\affil\{[^}]*\})', authors)

        for part in parts:
            if part.startswith('\\affil{'):
                # Estrai il numero dell'affiliazione
                match = re.search(r'\\affil\{([^}]*)\}', part)
                if match:
                    num = match.group(1)
                    run = authors_paragraph.add_run(num)
                    run.font.superscript = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
            else:
                run = authors_paragraph.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)

    def add_affiliations(self, affiliations):
        """Aggiunge le affiliazioni (10pt, centrato)"""
        affil_paragraph = self.doc.add_paragraph()
        affil_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        affil_paragraph.paragraph_format.space_after = Pt(12)
        affil_paragraph.paragraph_format.first_line_indent = Inches(0)

        # Gestisci gli apici
        parts = re.split(r'(\\affil\{[^}]*\})', affiliations)

        for part in parts:
            if part.startswith('\\affil{'):
                match = re.search(r'\\affil\{([^}]*)\}', part)
                if match:
                    num = match.group(1)
                    run = affil_paragraph.add_run(num)
                    run.font.superscript = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
            else:
                run = affil_paragraph.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

    def add_formatted_text(self, paragraph, text, font_size=11, bold=False):
        """Aggiunge testo con formattazione corsiva"""
        # Dividi per corsivi (pattern *testo*)
        parts = re.split(r'(\*[^*]+\*)', text)

        for part in parts:
            if part.startswith('*') and part.endswith('*'):
                # Testo in corsivo
                run = paragraph.add_run(part[1:-1])
                run.font.italic = True
            else:
                run = paragraph.add_run(part)

            run.font.name = 'Times New Roman'
            run.font.size = Pt(font_size)
            if bold:
                run.font.bold = True

    def add_body_text(self, body):
        """Aggiunge il corpo del testo con formattazione corretta"""
        # Dividi in paragrafi
        paragraphs = body.split('\n\n')

        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            paragraph = self.doc.add_paragraph()
            self.add_formatted_text(paragraph, para_text)

    def add_bibliography(self, entries):
        """Aggiunge la bibliografia"""
        if not entries:
            return

        # Titolo Bibliografia (maiuscoletto grassetto, centrato)
        bib_title = self.doc.add_paragraph()
        bib_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bib_title.paragraph_format.space_before = Pt(12)
        bib_title.paragraph_format.space_after = Pt(6)
        bib_title.paragraph_format.first_line_indent = Inches(0)

        title_run = bib_title.add_run("BIBLIOGRAFIA")
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(11)
        title_run.font.bold = True
        title_run.font.small_caps = True

        # Voci bibliografiche
        for entry in entries:
            if entry.strip():
                bib_paragraph = self.doc.add_paragraph()
                bib_paragraph.paragraph_format.first_line_indent = Inches(0)  # Nessuna indentazione per bibliografia

                self.add_formatted_text(bib_paragraph, entry.strip())

    def convert_latex_file(self, latex_file_path, output_path=None):
        """Converte un file LaTeX in DOCX"""
        with open(latex_file_path, 'r', encoding='utf-8') as f:
            latex_content = f.read()

        return self.convert_latex_string(latex_content, output_path)

    def convert_latex_string(self, latex_content, output_path=None):
        """Converte una stringa LaTeX in DOCX"""
        # Estrai contenuti
        content = self.extract_latex_content(latex_content)

        # Debug: stampa i contenuti estratti
        print("DEBUG - Contenuti estratti:")
        print(f"Titolo: {content['title']}")
        print(f"Autori: {content['authors']}")
        print(f"Affiliazioni: {content['affiliations']}")
        print(f"Corpo (primi 100 char): {content['body'][:100]}...")
        print(f"Bibliografia: {len(content['bibliography'])} voci")

        # Costruisci il documento
        if content['title']:
            self.add_title(content['title'])

        if content['authors']:
            self.add_authors(content['authors'])

        if content['affiliations']:
            self.add_affiliations(content['affiliations'])

        if content['body']:
            self.add_body_text(content['body'])

        if content['bibliography']:
            self.add_bibliography(content['bibliography'])

        # Salva il documento
        if output_path is None:
            output_path = 'gatm_abstract_fixed.docx'

        self.doc.save(output_path)
        return output_path

def main():
    """Funzione principale"""
    import argparse

    parser = argparse.ArgumentParser(description='Converte abstract LaTeX in formato DOCX GATM (versione corretta)')
    parser.add_argument('input_file', help='File LaTeX di input')
    parser.add_argument('-o', '--output', help='File DOCX di output')

    args = parser.parse_args()

    converter = FixedLatexToGATMConverter()
    output_file = converter.convert_latex_file(args.input_file, args.output)

    print(f"✅ Conversione completata: {output_file}")

if __name__ == "__main__":
    main()
